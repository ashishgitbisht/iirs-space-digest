# -*- coding: utf-8 -*-
"""
IIRS Space Digest - Combined Production Version
Generates:
1) HTML digest
2) DOCX digest
Daily automated space news for IIRS employees
LAST 24 HOURS ROLLING WINDOW
"""

# =========================
# Imports
# =========================
import os
import re
import html
import time
import email.utils
import requests
import feedparser

from io import BytesIO
from datetime import datetime, date, timedelta, timezone
from urllib.parse import urlparse, parse_qs, unquote, urljoin

from bs4 import BeautifulSoup
from newspaper import Article, Config
from googlenewsdecoder import gnewsdecoder

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


print("🚀 Starting IIRS Daily Space Digest - LAST 24 HOURS WINDOW...")


# =========================
# Filters and Feed Lists
# =========================

EXCLUDED_KEYWORDS = r'(?i)(rape|murder|KYC|digilocker|arrest|crime|FIR|strikes|Rajya Sabha|Muslims|metro)'

REGIONAL_KEYWORDS = r'(?i)(space|satellite|remote sensing|gis|iirs|rrsc|nrsc|earth observation|glacier|landslide|cloudburst|disaster|floods|avalanche|earthquake|seismic|hyperspectral|air quality index| AQI |snowfall)'

NATIONAL_KEYWORDS = r'(?i)(isro|nrsc|nsil|chandrayaan| IIST |gaganyaan|pslv|glsv|lvm3|spadex|gsat|insat|resourcesat|cartosat|risat|launch|rocket|spacecraft|astronaut|shukrayaan|aditya|spaceport|sriharikota|indian space|vyommitra|eos|pslv-c62|axiom|nesac|nsss|sslv|nvs|hlvm3|om1)'

INTERNATIONAL_KEYWORDS = r'(?i)(nasa|esa|jaxa|cnsa|roscosmos|spacex|blue origin|artemis|starship|crew dragon|iss|international space station|hubble|james webb|mars rover|perseverance|insight|booster|orbital|launch|spacecraft|astronaut|spacewalk|satellite|mission|space agency)'

REGIONAL_FEEDS = [
    'https://www.amarujala.com/rss/uttarakhand.rss',
    'https://khabardevbhoomi.com/feed/',
    'https://devbhoomimedia.com/feed',
    'https://pioneeredge.in/feed',
    'https://www.livehindustan.com/uttarakhand/rss',
    'https://timesofindia.indiatimes.com/city/delhi/rssfeeds/1311474.cms',
    'https://indianexpress.com/section/cities/delhi/feed/',
    'https://www.hindustantimes.com/cities/delhi-news/rssfeed/',
]

yesterday = (date.today() - timedelta(days=1)).strftime('%Y-%m-%d')
google_isro = f'https://news.google.com/rss/search?q=ISRO+OR+NRSC+OR+IIRS+after:{yesterday}&hl=en-IN&gl=IN&-site:indianexpress.com&-site:thehindu.com&-site:timesofindia.indiatimes.com&-site:isro.gov.in&-site:economictimes.indiatimes.com'

NATIONAL_FEEDS = [
    'https://timesofindia.indiatimes.com/rssfeeds/1201659.cms',
    'https://indianexpress.com/section/science/feed/',
    'https://www.thehindu.com/sci-tech/science/rssfeed/',
    'https://www.thehindu.com/news/national/rssfeed/',
    'https://www.isro.gov.in/rssnews.xml',
    'https://government.economictimes.indiatimes.com/rss/digital-india',
    'https://government.economictimes.indiatimes.com/rss/policy',
    'https://government.economictimes.indiatimes.com/rss/governance',
    'https://government.economictimes.indiatimes.com/rss/smart-infra',
    'https://government.economictimes.indiatimes.com/rss/Defence',
    'https://government.economictimes.indiatimes.com/rss/economy',
    google_isro
]

INTERNATIONAL_FEEDS = [
    'https://www.esa.int/rss/rss-topnews.xml',
    'https://www.esa.int/rss/programmes.xml',
    'https://www.esa.int/rss/space_science.xml',
    'https://www.esa.int/rss/earth_observation.xml',
    'https://www.nasa.gov/rss/dyn/breaking_news.rss',
    'https://www.nasa.gov/rss/dyn/images_of_the_day.rss',
    'https://www.space.com/feeds/all',
    'https://spaceflightnow.com/feed/',
    'https://phys.org/rss-feed/space-news/',
    'https://www.thespacereview.com/rss.xml',
    'https://interestingengineering.com/feed',
]


# =========================
# Image Extraction Helpers
# =========================

BAD_IMAGE_HINTS = [
    "logo", "icon", "favicon", "sprite", "banner", "ads", "advert",
    "google-news", "gnews", "default", "placeholder", "avatar",
    "feedburner", "newsletter", "branding", "youtube", "facebook",
    "twitter", "instagram", "linkedin", "whatsapp", "telegram",
    "share", "social", "theme-assets", "thumb", "thumbnail", "small"
]

BAD_IMAGE_EXTENSIONS = [".svg", ".ico"]


def is_valid_image_url(url):
    if not url or not url.startswith("http"):
        return False

    low = url.lower()

    if any(low.endswith(ext) for ext in BAD_IMAGE_EXTENSIONS):
        return False

    if any(hint in low for hint in BAD_IMAGE_HINTS):
        return False

    if "/wp-content/themes/" in low:
        return False

    return True


def normalize_img_url(img, base_url):
    if not img:
        return None
    img = img.strip().replace("\\/", "/")
    if img.startswith("//"):
        img = "https:" + img
    elif img.startswith("/"):
        img = urljoin(base_url, img)
    return img


def score_image_url(img_url):
    if not img_url:
        return -999

    score = 0
    low = img_url.lower()

    if any(x in low for x in ["og:image", "og-image"]):
        score += 30
    if any(x in low for x in ["hero", "featured", "lead", "main", "article"]):
        score += 20
    if any(x in low for x in ["thumb", "thumbnail", "small", "icon", "logo", "sprite"]):
        score -= 40
    if any(x in low for x in ["120x", "150x", "180x", "200x", "300x"]):
        score -= 25
    if any(ext in low for ext in [".jpg", ".jpeg", ".png", ".webp"]):
        score += 5

    return score


def resolve_google_news_url(url):
    if not url or "news.google.com" not in url:
        return url

    try:
        decoded = gnewsdecoder(url)
        if isinstance(decoded, dict) and decoded.get("status"):
            decoded_url = decoded.get("decoded_url")
            if decoded_url and decoded_url.startswith("http"):
                return decoded_url
    except:
        pass

    return url


def resolve_msn_original_url(url):
    if not url or 'msn.com' not in url:
        return url

    try:
        headers = {
            'User-Agent': 'Mozilla/5.0',
            'Accept-Language': 'en-US,en;q=0.9'
        }
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        html_text = response.text

        soup = BeautifulSoup(html_text, 'html.parser')

        canonical = soup.find('link', rel='canonical')
        if canonical and canonical.get('href'):
            canon_url = canonical['href'].strip()
            if canon_url.startswith('http') and 'msn.com' not in canon_url:
                return canon_url

        og_url = soup.find('meta', attrs={'property': 'og:url'})
        if og_url and og_url.get('content'):
            og_val = og_url['content'].strip()
            if og_val.startswith('http') and 'msn.com' not in og_val:
                return og_val

        for tag in soup.find_all('meta'):
            for attr in ['content', 'value']:
                val = tag.get(attr)
                if val and isinstance(val, str) and val.startswith('http'):
                    if 'msn.com' not in val and 'assets.msn.com' not in val and 'static.msn.com' not in val:
                        return val.strip()

        candidates = re.findall(r'https?://[^\s"\'<>\\]+', html_text)
        for cand in candidates:
            cand = unquote(cand.strip())
            if (
                cand.startswith('http')
                and 'msn.com' not in cand
                and 'assets.msn.com' not in cand
                and 'static.msn.com' not in cand
                and not any(x in cand.lower() for x in ['facebook.com', 'twitter.com', 'instagram.com', 'youtube.com'])
            ):
                return cand

        parsed = urlparse(url)
        query = parse_qs(parsed.query)
        for key in ['url', 'src', 'source', 'redirect', 'u']:
            if key in query:
                possible = unquote(query[key][0])
                if possible.startswith('http') and 'msn.com' not in possible:
                    return possible

    except Exception:
        pass

    return url


def resolve_final_article_url(url):
    if not url:
        return url

    if 'news.google.com' in url:
        url = resolve_google_news_url(url)

    if 'msn.com' in url:
        url = resolve_msn_original_url(url)

    return url


def extract_image_from_raw_html(url):
    try:
        headers = {
            "User-Agent": "Mozilla/5.0",
            "Accept-Language": "en-US,en;q=0.9"
        }
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        html_text = response.text

        patterns = [
            r'<meta[^>]+property=["\']og:image["\'][^>]+content=["\']([^"\']+)["\']',
            r'<meta[^>]+property=["\']og:image:url["\'][^>]+content=["\']([^"\']+)["\']',
            r'<meta[^>]+name=["\']twitter:image["\'][^>]+content=["\']([^"\']+)["\']',
            r'<meta[^>]+name=["\']twitter:image:src["\'][^>]+content=["\']([^"\']+)["\']',
            r'<img[^>]+data-lazy-src=["\']([^"\']+)["\']',
            r'<img[^>]+data-src=["\']([^"\']+)["\']',
            r'<img[^>]+data-srcset=["\']([^"\']+)["\']',
            r'<img[^>]+src=["\']([^"\']+)["\']'
        ]

        for pattern in patterns:
            matches = re.findall(pattern, html_text, flags=re.I)
            for match in matches:
                img = match.strip().split()[0]
                if img.startswith("//"):
                    img = "https:" + img
                elif img.startswith("/"):
                    img = urljoin(url, img)
                if is_valid_image_url(img):
                    return img
    except:
        pass

    return None


def extract_image_from_jsonld_or_scripts(url):
    try:
        headers = {
            "User-Agent": "Mozilla/5.0",
            "Accept-Language": "en-US,en;q=0.9"
        }
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        html_text = response.text

        patterns = [
            r'"image"\s*:\s*"([^"]+)"',
            r'"thumbnailUrl"\s*:\s*"([^"]+)"',
            r'"contentUrl"\s*:\s*"([^"]+)"',
            r'"url"\s*:\s*"([^"]+\.(?:jpg|jpeg|png|webp))"'
        ]

        for pattern in patterns:
            matches = re.findall(pattern, html_text, flags=re.I)
            for img in matches:
                img = img.replace("\\/", "/")
                if img.startswith("//"):
                    img = "https:" + img
                elif img.startswith("/"):
                    img = urljoin(url, img)
                if is_valid_image_url(img):
                    return img
    except:
        pass

    return None


def extract_image_with_newspaper(url):
    if not url or not url.startswith("http"):
        return None

    try:
        config = Config()
        config.browser_user_agent = 'Mozilla/5.0'
        config.request_timeout = 20

        article = Article(url, config=config)
        article.download()
        article.parse()

        if article.top_image and article.top_image.startswith("http") and is_valid_image_url(article.top_image):
            return article.top_image

        if article.images:
            images = []
            for img in article.images:
                if isinstance(img, str) and img.startswith("http") and is_valid_image_url(img):
                    images.append(img)
            if images:
                images = sorted(images, key=score_image_url, reverse=True)
                return images[0]
    except:
        pass

    return None


def extract_first_image_url(entry, article_url=None):
    article_url = resolve_final_article_url(article_url) if article_url else None

    if article_url:
        image = extract_image_with_newspaper(article_url)
        if image:
            return image

        image = extract_image_from_raw_html(article_url)
        if image:
            return image

        image = extract_image_from_jsonld_or_scripts(article_url)
        if image:
            return image

    try:
        for item in entry.get("media_content", []):
            url = item.get("url")
            if is_valid_image_url(url):
                return url
    except:
        pass

    try:
        for item in entry.get("media_thumbnail", []):
            url = item.get("url")
            if is_valid_image_url(url):
                return url
    except:
        pass

    try:
        for link in entry.get("links", []):
            href = link.get("href", "")
            link_type = link.get("type", "")
            rel = link.get("rel", "")
            if href and href.startswith("http") and (rel == "enclosure" or str(link_type).startswith("image/")):
                if is_valid_image_url(href):
                    return href
    except:
        pass

    return None


def try_download_image(image_url, timeout=20):
    if not image_url:
        return None

    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        response = requests.get(image_url, headers=headers, timeout=timeout)
        response.raise_for_status()
        content_type = response.headers.get("Content-Type", "").lower()

        if "image" not in content_type:
            return None

        return BytesIO(response.content)
    except Exception:
        return None


# =========================
# Text Helpers
# =========================

def clean_source_name(source_name):
    if not source_name:
        return "Unknown Source"
    source_name = html.unescape(str(source_name))
    return re.sub(r'\.\.\.$', '', source_name).strip()


def sanitize_filename(name):
    return re.sub(r'[^a-zA-Z0-9_-]+', '_', name)


def normalize_text(text):
    if not text:
        return ""

    text = html.unescape(str(text))

    replacements = {
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u00a0": " ",
        "\u200b": "",
        "\ufeff": "",
        "\\|": "|",
        "\\'": "'",
        '\\"': '"',
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\r\n?', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def clean_body_text(text, title=""):
    text = normalize_text(text)
    if not text:
        return ""

    lines = [ln.strip() for ln in text.splitlines()]

    bad_phrases = [
        "your browser does not support javascript",
        "related articles",
        "add asianet newsable as a preferred source",
        "google news",
        "follow us on",
        "read more",
        "advertisement",
        "recommended stories",
        "suggested articles",
        "share this article",
        "click here",
    ]

    cleaned = []
    seen = set()

    for ln in lines:
        if not ln:
            continue

        low = ln.lower().strip()

        if any(bp in low for bp in bad_phrases):
            continue

        if title and low == normalize_text(title).lower():
            continue

        if len(low) < 3:
            continue

        if low in seen:
            continue

        seen.add(low)
        cleaned.append(ln)

    text = "\n".join(cleaned)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text).strip()

    return text


def split_into_paragraphs(text):
    text = clean_body_text(text)
    if not text:
        return []

    paras = re.split(r'\n{2,}', text)
    final_paras = []

    for para in paras:
        para = re.sub(r'\s+', ' ', para).strip()
        if not para:
            continue
        if len(para) < 20:
            continue
        final_paras.append(para)

    return final_paras


def sanitize_html_content(text):
    if not text:
        return ''
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text[:380] + '...' if len(text) > 380 else text.strip()


# =========================
# News Timing
# =========================

def is_within_last_24_hours(entry):
    now = datetime.now(timezone.utc)
    cutoff_time = now - timedelta(hours=24)

    pub_struct = entry.get('published_parsed') or entry.get('updated_parsed') or entry.get('created_parsed')

    if pub_struct:
        try:
            pub_time = datetime(*pub_struct[:6], tzinfo=timezone.utc)
            return pub_time >= cutoff_time
        except:
            pass

    date_str = entry.get('published') or entry.get('updated') or entry.get('created')
    if date_str:
        try:
            parsed_tuple = email.utils.parsedate_tz(date_str)
            if parsed_tuple:
                ts = email.utils.mktime_tz(parsed_tuple)
                pub_time = datetime.fromtimestamp(ts, timezone.utc)
                return pub_time >= cutoff_time
        except:
            pass

    return False


# =========================
# Feed Fetching
# =========================

def fetch_news_from_feeds(feeds, max_articles=6):
    news = []

    for url in feeds:
        try:
            feed = feedparser.parse(url)
            print(f"📱 {feed.feed.get('title', 'Unknown')} - checking...")

            for entry in feed.entries[:15]:
                if not is_within_last_24_hours(entry):
                    continue

                title_lower = entry.title.lower()
                raw_summary = entry.get('summary', '') or entry.get('description', '')
                summary_lower = raw_summary.lower()
                full_text_check = title_lower + " " + summary_lower

                if url in REGIONAL_FEEDS:
                    keyword_pattern = REGIONAL_KEYWORDS
                elif url in NATIONAL_FEEDS:
                    keyword_pattern = NATIONAL_KEYWORDS
                else:
                    keyword_pattern = INTERNATIONAL_KEYWORDS

                if not re.search(keyword_pattern, title_lower):
                    continue

                if re.search(EXCLUDED_KEYWORDS, full_text_check):
                    print(f"🗑️ REMOVED (Excluded content): {entry.title[:40]}...")
                    continue

                original_link = entry.link
                final_link = resolve_final_article_url(original_link)

                image_url = extract_first_image_url(entry, final_link)
                summary = sanitize_html_content(raw_summary)
                title = re.sub(r'<[^>]+>', '', entry.title)

                news.append({
                    'title': title,
                    'link': final_link,
                    'source': feed.feed.get('title', 'Space News'),
                    'summary': summary,
                    'image': image_url
                })

                print(f"✅ NEW (24h): {title[:60]}...")
                print(f"🔗 Original link: {original_link}")
                print(f"🔗 Final link: {final_link}")
                print(f"🖼️ Image found: {image_url}")

                if len(news) >= max_articles:
                    break

            if len(news) >= max_articles:
                break

        except Exception as e:
            print(f"⚠️ Skip {url}: {e}")

    return news


# =========================
# HTML Generator
# =========================

def make_articles_html(news_list):
    html_out = ""

    for i, item in enumerate(news_list, 1):
        image_html = ''
        if item.get("image"):
            image_html = (
                f'<img src="{item["image"]}" alt="Space news image" '
                f'class="card-image" loading="lazy" '
                f'onerror="this.style.display=\'none\'">'
            )

        html_out += f'''
            <div class="news-card">
                <div class="card-content">
                    {image_html}
                    <div class="card-title">
                        <a href="{item["link"]}" target="_blank" rel="noopener noreferrer">{i}. {item["title"]}</a>
                    </div>
                    <div class="card-source">{item["source"]}</div>
                    <div class="card-summary">{item["summary"]}</div>
                    <a href="{item["link"]}" target="_blank" rel="noopener noreferrer" class="read-more">Read Full Article →</a>
                </div>
            </div>
        '''

    return html_out


# =========================
# DOCX Helpers
# =========================

def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), 'A6A6A6')
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_top_border(paragraph, color="D9D9D9", size="6", space="4"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), size)
    top.set(qn('w:space'), space)
    top.set(qn('w:color'), color)
    pbdr.append(top)
    pPr.append(pbdr)


def add_box_border(paragraph, color="808080", size="8", space="8"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for side_name in ['top', 'left', 'bottom', 'right']:
        side = OxmlElement(f'w:{side_name}')
        side.set(qn('w:val'), 'single')
        side.set(qn('w:sz'), size)
        side.set(qn('w:space'), space)
        side.set(qn('w:color'), color)
        pbdr.append(side)
    pPr.append(pbdr)


def add_hyperlink(paragraph, text, url, color="0000FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single" if underline else "none")
    rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def set_section_columns(section, num_cols=1, space=360):
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)

    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), str(space))


def add_footer_to_section(section):
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)

    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""

    add_top_border(paragraph, color="D9D9D9", size="6", space="4")

    run = paragraph.add_run("Indian Institute of Remote Sensing (ISRO), Dehradun - 248001")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)


def apply_footer_to_all_sections(doc):
    for section in doc.sections:
        add_footer_to_section(section)


def fetch_full_article_text(url, fallback_summary="", title=""):
    fallback_summary = clean_body_text(fallback_summary, title=title)

    if not url or url == '#':
        return fallback_summary

    url = resolve_final_article_url(url)

    try:
        config = Config()
        config.browser_user_agent = 'Mozilla/5.0'
        config.request_timeout = 20

        article = Article(url, config=config)
        article.download()
        article.parse()

        text = clean_body_text(article.text or '', title=title)
        if len(text) >= 300:
            return text
    except Exception:
        pass

    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        raw_html = response.text

        raw_html = re.sub(r'<script.*?>.*?</script>', ' ', raw_html, flags=re.I | re.S)
        raw_html = re.sub(r'<style.*?>.*?</style>', ' ', raw_html, flags=re.I | re.S)

        patterns = [
            r'<article[^>]*>(.*?)</article>',
            r'<main[^>]*>(.*?)</main>',
            r'<div[^>]+class=["\'][^"\']*(?:article|story|content|main-content|post-content|entry-content|td-post-content|news-detail|story-detail)[^"\']*["\'][^>]*>(.*?)</div>'
        ]

        extracted = ''
        for pattern in patterns:
            matches = re.findall(pattern, raw_html, flags=re.I | re.S)
            if matches:
                flat = []
                for m in matches[:2]:
                    if isinstance(m, tuple):
                        flat.extend([x for x in m if x])
                    else:
                        flat.append(m)
                extracted = ' '.join(flat)
                break

        if not extracted:
            extracted = raw_html

        extracted = re.sub(r'</p>|<br\s*/?>|</div>|</section>|</article>|</li>|</h[1-6]>', '\n', extracted, flags=re.I)
        extracted = re.sub(r'<li[^>]*>', '- ', extracted, flags=re.I)
        extracted = re.sub(r'<[^>]+>', ' ', extracted)

        extracted = clean_body_text(extracted, title=title)

        if len(extracted) >= 300:
            return extracted
    except Exception:
        pass

    return fallback_summary


def add_article_body_in_two_columns(doc, paragraphs):
    if not paragraphs:
        paragraphs = ['Summary not available.']

    col_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_section_columns(col_section, num_cols=2, space=360)
    add_footer_to_section(col_section)

    for para in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        para = re.sub(r'\s+', ' ', para).strip()

        run = p.add_run(para)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)

    back_to_one = doc.add_section(WD_SECTION.CONTINUOUS)
    set_section_columns(back_to_one, num_cols=1, space=360)
    add_footer_to_section(back_to_one)


def generate_docx(news_items, output_path, digest_date_str):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    set_section_columns(section, num_cols=1)

    styles = doc.styles
    styles['Normal'].font.name = 'Times New Roman'
    styles['Normal'].font.size = Pt(11)

    add_footer_to_section(section)

    header_box = doc.add_paragraph()
    header_box.paragraph_format.space_after = Pt(10)

    tab_stops = header_box.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Inches(6.9), WD_TAB_ALIGNMENT.RIGHT)

    header_box.add_run("\t")

    run1 = header_box.add_run("IIRS Daily Space Digest")
    run1.bold = True
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(12)

    header_box.add_run("\t")

    run2 = header_box.add_run(digest_date_str)
    run2.bold = True
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(10)

    add_box_border(header_box, color="808080", size="8", space="8")

    doc.add_paragraph('')

    for idx, item in enumerate(news_items, start=1):
        title = normalize_text(item.get('title', 'Untitled'))
        source = clean_source_name(item.get('source', ''))
        link = resolve_final_article_url(normalize_text(item.get('link', '')))
        summary = normalize_text(item.get('summary', ''))
        image_url = item.get('image')

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'{idx}. {title}')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

        meta_parts = []
        if source:
            meta_parts.append(source)

        if meta_parts:
            meta = doc.add_paragraph()
            meta.paragraph_format.space_after = Pt(3)
            meta_run = meta.add_run(' | '.join(meta_parts))
            meta_run.italic = True
            meta_run.font.name = 'Times New Roman'
            meta_run.font.size = Pt(10)

        if link and link != '#':
            link_p = doc.add_paragraph()
            link_p.paragraph_format.space_after = Pt(4)
            link_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_hyperlink(link_p, "Read more", link)

        image_stream = try_download_image(image_url)
        if image_stream:
            try:
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_run = img_p.add_run()
                img_run.add_picture(image_stream, width=Inches(4.8))
                img_p.paragraph_format.space_after = Pt(6)
            except Exception:
                pass

        body_text = fetch_full_article_text(
            url=link,
            fallback_summary=summary,
            title=title
        )

        body_text = clean_body_text(body_text, title=title)
        body_paragraphs = split_into_paragraphs(body_text)

        if not body_paragraphs:
            fallback_clean = clean_body_text(summary, title=title)
            body_paragraphs = split_into_paragraphs(fallback_clean)

        add_article_body_in_two_columns(doc, body_paragraphs)

        if idx != len(news_items):
            sep = doc.add_paragraph()
            add_bottom_border(sep)
            doc.add_paragraph('')

    apply_footer_to_all_sections(doc)
    doc.save(output_path)
    print(f'DOCX saved: {output_path}')


# =========================
# Main Fetch
# =========================

print("🏔️ Fetching REGIONAL...")
regional_news = fetch_news_from_feeds(REGIONAL_FEEDS, max_articles=5)

print("🇮🇳 Fetching NATIONAL...")
national_news = fetch_news_from_feeds(NATIONAL_FEEDS, max_articles=6)

print("🌌 Fetching INTERNATIONAL...")
international_news = fetch_news_from_feeds(INTERNATIONAL_FEEDS, max_articles=8)

all_news = []
for news_list, category in [
    (regional_news, "🏔️ Regional Updates"),
    (national_news, "🇮🇳 National Updates"),
    (international_news, "🌌 International Updates")
]:
    for item in news_list:
        item['category'] = category
        all_news.append(item)

if not all_news:
    all_news.append({
        'title': 'No space news in last 24h',
        'link': '#',
        'source': 'IIRS Digest',
        'summary': 'Check back tomorrow!',
        'image': None,
        'category': 'System'
    })


# =========================
# HTML Output
# =========================

all_articles_html = make_articles_html(all_news)

ist_offset = timezone(timedelta(hours=5, minutes=30))
timestamp = datetime.now(ist_offset).strftime("%d-%m-%Y | %H:%M IST")

html_body = f"""<!DOCTYPE html>
<html data-theme="dark">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>
:root {{
    --bg-primary: #0a0a0a;
    --bg-secondary: rgba(10, 10, 10, 0.9);
    --card-bg: rgba(255,255,255,0.05);
    --card-summary: rgba(255,255,255,0.02);
    --text-primary: #c0c0c0;
    --text-secondary: #a0a0a0;
    --text-light: #d0d0d0;
    --text-white: #ffffff;
    --border-light: rgba(255,255,255,0.08);
    --border-card: rgba(255,255,255,0.1);
    --shadow-dark: rgba(0,0,0,0.8);
    --cyan-accent: #00ffff;
}}
[data-theme="light"] {{
    --bg-primary: #f8fafc !important;
    --bg-secondary: rgba(255, 255, 255, 0.98) !important;
    --card-bg: rgba(255,255,255,0.95) !important;
    --card-summary: rgba(248, 250, 252, 0.8) !important;
    --text-primary: #1e293b !important;
    --text-secondary: #475569 !important;
    --text-light: #334155 !important;
    --text-white: #0f172a !important;
    --border-light: rgba(0,0,0,0.06) !important;
    --border-card: rgba(0,0,0,0.08) !important;
    --shadow-dark: rgba(0,0,0,0.1) !important;
    --cyan-accent: #00b8d4 !important;
}}

* {{ box-sizing: border-box !important; }}
html {{ background: var(--bg-primary) !important; min-height: 100vh !important; }}
body {{
    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif !important;
    margin: 0 !important;
    padding: 20px !important;
    background: var(--bg-primary) !important;
    color: var(--text-primary) !important;
    min-height: 100vh !important;
    display: flex !important;
    flex-direction: column !important;
    align-items: center !important;
}}

body::before {{
    content: '' !important;
    position: fixed !important;
    top: 0; left: 0; width: 100%; height: 100%;
    background-image:
        radial-gradient(1px 1px at 20px 30px, rgba(255,255,255,0.4), transparent),
        radial-gradient(1px 1px at 160px 30px, rgba(255,255,255,0.25), transparent);
    background-size: 300px 300px !important;
    animation: voidDrift 60s linear infinite !important;
    pointer-events: none !important;
    z-index: -1 !important;
    opacity: 0.5 !important;
}}
@keyframes voidDrift {{ from {{ background-position: 0 0; }} to {{ background-position: 0 600px; }} }}

.theme-toggle {{
    position: fixed !important; top: 20px !important; right: 20px !important;
    width: 45px !important; height: 45px !important;
    border-radius: 50% !important; border: none !important;
    background: rgba(255,255,255,0.1) !important;
    color: #fff !important; font-size: 20px !important;
    cursor: pointer !important; backdrop-filter: blur(10px) !important;
    z-index: 1000 !important;
}}

.scroll-container {{
    width: 80% !important;
    max-width: none !important;
    min-width: 600px !important;
    background: var(--bg-secondary) !important;
    backdrop-filter: blur(30px) !important;
    border: 1px solid var(--border-light) !important;
    border-radius: 24px !important;
    padding: 40px !important;
    box-shadow: 0 35px 70px var(--shadow-dark) !important;
    margin-top: 20px !important;
}}

h2 {{
    color: var(--text-white) !important;
    text-align: center !important;
    border-bottom: 2px solid var(--border-light) !important;
    padding-bottom: 20px !important;
    margin-bottom: 30px !important;
    font-weight: 700 !important;
    letter-spacing: 1px !important;
}}

.news-card {{ margin-bottom: 40px !important; }}
.card-content {{
    background: var(--card-bg) !important;
    border: 1px solid var(--border-card) !important;
    border-radius: 20px !important;
    padding: 30px !important;
    box-shadow: 0 10px 30px var(--shadow-dark) !important;
    transition: transform 0.3s ease !important;
}}
.card-content:hover {{ transform: translateY(-5px) !important; border-color: var(--cyan-accent) !important; }}

.card-image {{
    width: 100% !important; height: 350px !important;
    object-fit: cover !important;
    border-radius: 12px !important; margin-bottom: 20px !important;
    border: 1px solid var(--border-card) !important;
}}

.card-title a {{
    color: var(--text-white) !important; text-decoration: none !important;
    font-size: 24px !important;
    font-weight: 600 !important; display: block !important;
    margin-bottom: 10px !important;
}}
.card-title a:hover {{ text-decoration: underline !important; color: var(--cyan-accent) !important; }}

.card-source {{
    display: inline-block !important; padding: 5px 12px !important;
    background: rgba(255,255,255,0.05) !important; border-radius: 15px !important;
    font-size: 13px !important; color: var(--text-secondary) !important;
    margin-bottom: 15px !important; border: 1px solid var(--border-light) !important;
}}

.card-summary {{
    color: var(--text-light) !important; line-height: 1.7 !important;
    font-size: 16px !important;
    margin-bottom: 20px !important;
}}

.read-more {{
    display: inline-block !important; padding: 10px 20px !important;
    background: transparent !important; border: 1px solid var(--cyan-accent) !important;
    color: var(--cyan-accent) !important; text-decoration: none !important;
    border-radius: 25px !important; font-weight: 600 !important; font-size: 14px !important;
    transition: all 0.3s ease !important;
}}
.read-more:hover {{ background: var(--cyan-accent) !important; color: #000 !important; }}

.footer {{
    text-align: center !important; margin-top: 40px !important;
    color: var(--text-secondary) !important; font-size: 13px !important;
    padding-bottom: 20px !important;
}}

@media (max-width: 1000px) {{
    .scroll-container {{ width: 90% !important; min-width: 0 !important; }}
}}
@media (max-width: 768px) {{
    .scroll-container {{ width: 95% !important; padding: 20px !important; }}
    .card-content {{ padding: 20px !important; }}
    h2 {{ font-size: 22px !important; }}
    .card-image {{ height: 200px !important; }}
}}
</style>
</head>
<body>
<button class="theme-toggle" id="themeToggle" title="Toggle Theme">☀️</button>

<div class="scroll-container">
    <h2>🌌 IIRS Daily Space Digest</h2>
    <p style="text-align:center; color:var(--text-secondary); margin-top:-20px; margin-bottom:40px;">
        {timestamp} | {len(all_news)} Updates Found
    </p>

    {all_articles_html}

    <div class="footer">
        IIRS Library | Indian Institute of Remote Sensing | Dehradun<br>
        <small>Automated Digest System</small>
    </div>
</div>

<script>
const btn = document.getElementById('themeToggle');
const html = document.documentElement;

if (localStorage.getItem('theme') === 'light') {{
    html.setAttribute('data-theme', 'light');
    btn.textContent = '🌙';
}}

btn.addEventListener('click', () => {{
    if (html.getAttribute('data-theme') === 'light') {{
        html.removeAttribute('data-theme');
        btn.textContent = '☀️';
        localStorage.setItem('theme', 'dark');
    }} else {{
        html.setAttribute('data-theme', 'light');
        btn.textContent = '🌙';
        localStorage.setItem('theme', 'light');
    }}
}});
</script>
</body>
</html>
"""

html_filename = f'IIRS_SpaceNews_Daily_{datetime.now().strftime("%Y%m%d")}.html'
with open(html_filename, 'w', encoding='utf-8') as f:
    f.write(html_body)

print(f"✅ SAVED: {html_filename} with {len(all_news)} items")


# =========================
# DOCX Output
# =========================

digest_date_str = datetime.now(ist_offset).strftime('%A, %d/%m/%Y')
docx_filename = f"iirs_daily_space_digest_{datetime.now(ist_offset).strftime('%d_%m_%Y')}.docx"

generate_docx(
    news_items=all_news,
    output_path=docx_filename,
    digest_date_str=digest_date_str
)

print("📱 HTML + DOCX generation complete.")
